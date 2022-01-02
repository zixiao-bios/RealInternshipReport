from docx import Document
from docx.document import Document as Doc
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Pt

if __name__ == '__main__':
    doc: Doc
    doc = Document()
    styles = doc.styles
    style = styles.add_style("表头", WD_STYLE_TYPE.PARAGRAPH)
    style.hidden = False
    style.quick_style = True
    style.priority = 1

    doc.add_paragraph('实习报告                                                     2021年1月2日', style="表头")
    table = doc.add_table(1, 1, "TableGrid")
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    table.rows[0].height = Pt(600)
    table.cell(0, 0).text = "这是正文"
    doc.add_page_break()

    doc.add_paragraph('实习报告                                                     2021年1月2日', style="表头")
    table = doc.add_table(1, 1, "TableGrid")
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    table.rows[0].height = Pt(600)
    table.cell(0, 0).text = "这是正文"
    doc.add_page_break()

    doc.save("test.docx")
