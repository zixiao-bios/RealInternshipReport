import random
from requests_html import HTMLSession, HTML

from sumy.parsers.html import HtmlParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer as Summarizer
from sumy.nlp.stemmers import Stemmer
from sumy.utils import get_stop_words

from docx import Document
from docx.document import Document as Doc
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Pt
from docx.oxml.ns import qn

generate_page_num = 5
year = 2021
month = 9
date = 27
max_date = 30

start_url = "https://www.runoob.com/"
LANGUAGE = "chinese"
SENTENCES_COUNT = 50


def refresh_datetime():
    global date, month, max_date
    date += 1
    if date > max_date:
        date = 1
        month += 1


def generate_paragraph(url):
    parser = HtmlParser.from_url(url, Tokenizer(LANGUAGE))
    # or for plain text files
    # parser = PlaintextParser.from_file("document.txt", Tokenizer(LANGUAGE))
    # parser = PlaintextParser.from_string(text, Tokenizer(LANGUAGE))
    stemmer = Stemmer(LANGUAGE)

    summarizer = Summarizer(stemmer)
    summarizer.stop_words = get_stop_words(LANGUAGE)

    paragraph_list = []
    build_str = "\t"
    for sentence in summarizer(parser.document, SENTENCES_COUNT):
        sentence = str(sentence).replace("尝试一下 » ", "")
        if random.random() < 0.5:
            paragraph_list.append(build_str)
            build_str = "\t"
        build_str += sentence
    return paragraph_list


def get_next_url(last_page_url):
    session = HTMLSession()
    r = session.get(last_page_url)
    html = HTML(html=r.text)
    next_link = html.find(".next-design-link", first=True)
    if len(next_link.absolute_links) == 0:
        return None
    print(next_link.text, next_link.absolute_links)
    for next_url in next_link.absolute_links:
        return next_url


def get_topic_url_set(home_url):
    url_set = set()
    session = HTMLSession()
    r = session.get(home_url)
    html = HTML(html=r.text)
    item_list = html.find(".item-1")
    for item in item_list:
        for url in item.absolute_links:
            url_set.add(url)
    return url_set


if __name__ == '__main__':
    # init doc
    doc: Doc
    doc = Document()
    styles = doc.styles
    style = styles.add_style("表头", WD_STYLE_TYPE.PARAGRAPH)
    style.hidden = False
    style.quick_style = True
    style.priority = 1

    print("正在爬取网页：")
    topic_url_set = get_topic_url_set(start_url)
    page_url = get_next_url(topic_url_set.pop())
    for i in range(generate_page_num):
        # input datetime
        run = doc.add_paragraph(style="表头").add_run()
        font = run.font
        font.name = "宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.add_text("实习记录                                                     "
                     + str(year) + "年"
                     + str(month) + "月"
                     + str(date) + "日")
        refresh_datetime()

        # input table text
        table = doc.add_table(1, 1, "TableGrid")
        table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        table.rows[0].height = Pt(600)
        text_len = 0
        while text_len < 400:
            while page_url is None:
                # find another topic
                page_url = get_next_url(topic_url_set.pop())
            for paragraph in generate_paragraph(page_url):
                text_len += len(paragraph)
                table.cell(0, 0).add_paragraph(paragraph)
                if text_len > 600:
                    break
            # goto next webpage
            page_url = get_next_url(page_url)
            if text_len > 600:
                break

        # move to next page
        doc.add_page_break()
    doc.save("real internship report.docx")
